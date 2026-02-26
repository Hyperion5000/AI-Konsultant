import os
import sys
import glob
import argparse
import logging
import pickle
from typing import List, Optional
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_community.retrievers import BM25Retriever
from langchain_core.documents import Document
import docx

# Ensure we can import bot.config
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

# Import configuration
try:
    from bot import config
    from bot.core.logger import get_logger, configure_logging
except ImportError:
    # If run from root without installation, this might fail unless pythonpath is set.
    # We'll try to load from env or use defaults if config import fails.
    load_dotenv()
    class Config:
        EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "cointegrated/rubert-tiny2")
        CHROMA_DIR = os.getenv("CHROMA_DIR", "storage/chroma")
        BM25_DIR = os.getenv("BM25_DIR", "storage/bm25")
        STORAGE_DIR = "storage"
    config = Config()

    # Fallback logging
    def configure_logging():
        logging.basicConfig(level=logging.INFO)
    def get_logger(name):
        return logging.getLogger(name)

# Configure logging
configure_logging()
logger = get_logger(__name__)

def load_documents(data_dir: str) -> List[Document]:
    """
    Scans the data directory for .docx, .txt, .md, and .pdf files and loads their content.
    """
    documents = []

    # 1. Load .docx files
    docx_files = glob.glob(os.path.join(data_dir, "**", "*.docx"), recursive=True)
    if docx_files:
        logger.info(f"Found {len(docx_files)} .docx files.")
        for file_path in docx_files:
            try:
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                text = "\n".join(full_text)
                if text.strip():
                    filename = os.path.basename(file_path)
                    documents.append(Document(page_content=text, metadata={"source": filename, "type": "docx"}))
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")

    # 2. Load .txt and .md files
    text_files = glob.glob(os.path.join(data_dir, "**", "*.txt"), recursive=True) + \
                 glob.glob(os.path.join(data_dir, "**", "*.md"), recursive=True)

    if text_files:
        logger.info(f"Found {len(text_files)} text/markdown files.")
        for file_path in text_files:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                if text.strip():
                    filename = os.path.basename(file_path)
                    ext = os.path.splitext(filename)[1].lower().replace(".", "")
                    documents.append(Document(page_content=text, metadata={"source": filename, "type": ext}))
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")

    # 3. Load .pdf files
    pdf_files = glob.glob(os.path.join(data_dir, "**", "*.pdf"), recursive=True)
    if pdf_files:
        logger.info(f"Found {len(pdf_files)} .pdf files.")
        for file_path in pdf_files:
            try:
                import fitz  # pymupdf
                doc = fitz.open(file_path)
                full_text = []
                for page in doc:
                    full_text.append(page.get_text())
                text = "\n".join(full_text)
                if text.strip():
                    filename = os.path.basename(file_path)
                    documents.append(Document(page_content=text, metadata={"source": filename, "type": "pdf"}))
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")

    logger.info(f"Total documents loaded: {len(documents)}")
    return documents

def get_text_splitter(chunk_size: int, chunk_overlap: int) -> RecursiveCharacterTextSplitter:
    """
    Returns a configured text splitter for legal documents.
    """
    # Smart chunking for legal texts
    # separators order matters: try to split by article/chapter first, then paragraphs, then sentences
    separators = ["\nСтатья ", "\nГлава ", "\n\n", "\n", " ", ""]

    return RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators,
        is_separator_regex=False
    )

def create_vector_store(
    documents: List[Document],
    chroma_dir: str,
    bm25_dir: str,
    embedding_model_name: str,
    chunk_size: int,
    chunk_overlap: int
):
    """
    Splits documents, adds metadata, and creates a Chroma vector store + BM25 index.
    """
    if not documents:
        logger.warning("No documents to index.")
        return

    text_splitter = get_text_splitter(chunk_size, chunk_overlap)
    split_docs = text_splitter.split_documents(documents)

    # Add chunk_id and title (optional) to metadata
    for i, doc in enumerate(split_docs):
        doc.metadata["chunk_id"] = i
        # Attempt to extract a title from the first line if likely
        # This is a heuristic; 'source' is the main identifier.
        lines = doc.page_content.splitlines()
        if lines:
            possible_title = lines[0].strip()[:100] # Cap title length
            doc.metadata["title"] = possible_title

    logger.info(f"Split into {len(split_docs)} chunks.")

    # 1. Create and save BM25 Retriever
    logger.info("Creating BM25 retriever...")
    try:
        bm25_retriever = BM25Retriever.from_documents(split_docs)
        # We save the retriever using pickle
        if not os.path.exists(bm25_dir):
            os.makedirs(bm25_dir)

        bm25_path = os.path.join(bm25_dir, "bm25_retriever.pkl")
        with open(bm25_path, "wb") as f:
            pickle.dump(bm25_retriever, f)
        logger.info(f"BM25 retriever saved to {bm25_path}")
    except Exception as e:
        logger.error(f"Failed to create BM25 retriever: {e}")

    # 2. Create Chroma DB
    logger.info(f"Initializing embeddings model: {embedding_model_name}...")
    try:
        embeddings = HuggingFaceEmbeddings(model_name=embedding_model_name)
    except Exception as e:
        logger.error(f"Failed to initialize embeddings: {e}")
        return

    logger.info(f"Creating Chroma vector store in {chroma_dir}...")
    try:
        if os.path.exists(chroma_dir) and os.path.exists(os.path.join(chroma_dir, "chroma.sqlite3")):
             # Check if we should append or warn. Chroma default is append.
             logger.warning(f"Chroma DB already exists in {chroma_dir}. Appending.")

        vector_store = Chroma.from_documents(
            documents=split_docs,
            embedding=embeddings,
            persist_directory=chroma_dir
        )
        logger.info(f"Vector store created and persisted to {chroma_dir}.")
    except Exception as e:
        logger.error(f"Failed to create vector store: {e}")

def main():
    parser = argparse.ArgumentParser(description="Create RAG index from documents.")

    # Use config defaults
    default_input = os.path.join(getattr(config, 'STORAGE_DIR', 'storage'), 'documents')

    parser.add_argument("--input", default=default_input, help="Directory containing documents (.docx, .txt, .md)")
    parser.add_argument("--persist-dir", default=config.CHROMA_DIR, help="Directory to store Chroma DB")
    parser.add_argument("--bm25-dir", default=config.BM25_DIR, help="Directory to store BM25 index")
    parser.add_argument("--embedding-model", default=config.EMBEDDING_MODEL, help="HuggingFace embedding model name")
    parser.add_argument("--chunk-size", type=int, default=1000, help="Chunk size for splitting text")
    parser.add_argument("--chunk-overlap", type=int, default=200, help="Chunk overlap size")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        logger.info(f"Input directory '{args.input}' does not exist. Creating it.")
        os.makedirs(args.input, exist_ok=True)
        logger.info(f"Please put your documents (.docx, .txt, .md) in '{args.input}' and run this script again.")
        return

    documents = load_documents(args.input)
    if documents:
        create_vector_store(
            documents,
            chroma_dir=args.persist_dir,
            bm25_dir=args.bm25_dir,
            embedding_model_name=args.embedding_model,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap
        )
    else:
        logger.warning(f"No documents found in {args.input}.")

if __name__ == "__main__":
    main()
